#!/usr/bin/env python3
"""Render a lightweight markup file into a styled institutional research DOCX."""
import re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY   = RGBColor(0x0F, 0x2B, 0x46)
ACCENT = RGBColor(0xA8, 0x6B, 0x1E)
GREY   = RGBColor(0x55, 0x5F, 0x6B)
BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
HDR_BG = "0F2B46"
ALT_BG = "F2F4F7"
RULE   = "C8CDD4"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:color'), 'auto'); el.set(qn('w:fill'), hexcolor)
    tcPr.append(el)


def cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('start', left), ('bottom', bottom), ('end', right)):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); mar.append(e)
    tcPr.append(mar)


def table_borders(table, colour=RULE, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz))
        e.set(qn('w:space'), '0'); e.set(qn('w:color'), colour)
        borders.append(e)
    tblPr.append(borders)


def para_rule(p, colour="0F2B46", sz=8):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:bottom')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '4'); b.set(qn('w:color'), colour)
    pbdr.append(b); pPr.append(pbdr)


def left_bar(p, colour="A86B1E", sz=18):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    b = OxmlElement('w:left')
    b.set(qn('w:val'), 'single'); b.set(qn('w:sz'), str(sz))
    b.set(qn('w:space'), '10'); b.set(qn('w:color'), colour)
    pbdr.append(b); pPr.append(pbdr)


def shade_para(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    el = OxmlElement('w:shd')
    el.set(qn('w:val'), 'clear'); el.set(qn('w:color'), 'auto'); el.set(qn('w:fill'), hexcolor)
    pPr.append(el)



# --- WordprocessingML child-order normalisation -------------------------------
# Word rejects files whose w:pPr / w:tblPr / w:tcPr children are out of schema
# order. python-docx appends our custom elements at the end, so we reorder the
# whole tree once, immediately before saving.

_ORDER = {
    'pPr': ['pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
            'widowControl', 'numPr', 'suppressLineNumbers', 'pBdr', 'shd', 'tabs',
            'suppressAutoHyphens', 'kinsoku', 'wordWrap', 'overflowPunct',
            'topLinePunct', 'autoSpaceDE', 'autoSpaceDN', 'bidi', 'adjustRightInd',
            'snapToGrid', 'spacing', 'ind', 'contextualSpacing', 'mirrorIndents',
            'suppressOverlap', 'jc', 'textDirection', 'textAlignment',
            'textboxTightWrap', 'outlineLvl', 'divId', 'cnfStyle', 'rPr', 'sectPr',
            'pPrChange'],
    'tblPr': ['tblStyle', 'tblpPr', 'tblOverlap', 'bidiVisual', 'tblStyleRowBandSize',
              'tblStyleColBandSize', 'tblW', 'jc', 'tblCellSpacing', 'tblInd',
              'tblBorders', 'shd', 'tblLayout', 'tblCellMar', 'tblLook', 'tblCaption',
              'tblDescription', 'tblPrChange'],
    'tcPr': ['cnfStyle', 'tcW', 'gridSpan', 'hMerge', 'vMerge', 'tcBorders', 'shd',
             'noWrap', 'tcMar', 'textDirection', 'tcFitText', 'vAlign', 'hideMark',
             'headers', 'cellIns', 'cellDel', 'cellMerge', 'tcPrChange'],
    'rPr': ['rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps', 'strike',
            'dstrike', 'outline', 'shadow', 'emboss', 'imprint', 'noProof',
            'snapToGrid', 'vanish', 'webHidden', 'color', 'spacing', 'w', 'kern',
            'position', 'sz', 'szCs', 'highlight', 'u', 'effect', 'bdr', 'shd',
            'fitText', 'vertAlign', 'rtl', 'cs', 'em', 'lang', 'eastAsianLayout',
            'specVanish', 'oMath', 'rPrChange'],
}
_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'


def normalise_order(document):
    for parent in document.element.body.iter():
        tag = parent.tag
        if not tag.startswith(_W):
            continue
        local = tag[len(_W):]
        order = _ORDER.get(local)
        if not order:
            continue
        kids = list(parent)
        idx = {name: n for n, name in enumerate(order)}
        def key(el, _c=[0]):
            _c[0] += 1
            name = el.tag[len(_W):] if el.tag.startswith(_W) else ''
            return (idx.get(name, len(order) + 1), _c[0])
        for el in sorted(kids, key=key):
            parent.append(el)


INLINE = re.compile(r'(\*\*.+?\*\*|\*[^*]+?\*)')


def add_runs(p, text, size=9.5, colour=BLACK, base_bold=False, font='Georgia'):
    for tok in INLINE.split(text):
        if not tok:
            continue
        bold, italic = base_bold, False
        t = tok
        if tok.startswith('**') and tok.endswith('**') and len(tok) > 4:
            bold, t = True, tok[2:-2]
        elif tok.startswith('*') and tok.endswith('*') and len(tok) > 2:
            italic, t = True, tok[1:-1]
        r = p.add_run(t)
        r.font.size = Pt(size); r.font.color.rgb = colour
        r.font.name = font; r.bold = bold; r.italic = italic
        r._element.rPr.rFonts.set(qn('w:eastAsia'), font)


FOOTER = 'Liontrust Asset Management PLC (LIO LN) — Initiation of Coverage  ·  30 August 2026'


def build(src_path, out_path):
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Georgia'; st.font.size = Pt(9.5)
    st.element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
    st.paragraph_format.space_after = Pt(5)
    st.paragraph_format.line_spacing = 1.14

    sec = doc.sections[0]
    sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(1.9); sec.right_margin = Cm(1.9)

    # running footer
    ftr = sec.footer.paragraphs[0]
    ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = ftr.add_run(FOOTER + '  ·  For discussion purposes only — not investment advice')
    fr.font.size = Pt(6.5); fr.font.color.rgb = GREY; fr.font.name = 'Georgia'

    lines = open(src_path, encoding='utf-8').read().split('\n')
    i = 0
    while i < len(lines):
        ln = lines[i]; s = ln.strip()

        if s == '[PAGEBREAK]':
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE); i += 1; continue
        if not s:
            i += 1; continue

        # ---- tables ----
        if s.startswith('|'):
            block = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                block.append(lines[i].strip()); i += 1
            rows = [[c.strip() for c in r.strip('|').split('|')] for r in block
                    if not re.match(r'^\|[\s\-:|]+\|$', r)]
            ncol = max(len(r) for r in rows)
            rows = [r + [''] * (ncol - len(r)) for r in rows]
            tbl = doc.add_table(rows=len(rows), cols=ncol)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            tbl.autofit = True
            table_borders(tbl)
            fs = 8.0 if ncol <= 6 else (7.2 if ncol <= 8 else 6.6)
            for ri, row in enumerate(rows):
                for ci, val in enumerate(row):
                    cell = tbl.cell(ri, ci)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_after = Pt(1.5)
                    p.paragraph_format.space_before = Pt(1.5)
                    p.paragraph_format.line_spacing = 1.0
                    cell_margins(cell)
                    if ri == 0:
                        shade(cell, HDR_BG)
                        add_runs(p, val, size=fs, colour=RGBColor(0xFF, 0xFF, 0xFF), base_bold=True)
                    else:
                        if ri % 2 == 0:
                            shade(cell, ALT_BG)
                        add_runs(p, val, size=fs, colour=BLACK)
            doc.add_paragraph().paragraph_format.space_after = Pt(3)
            continue

        # ---- headings ----
        if s.startswith('#### '):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
            add_runs(p, s[5:], size=9.5, colour=ACCENT, base_bold=True); i += 1; continue
        if s.startswith('### '):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
            add_runs(p, s[4:], size=10.5, colour=NAVY, base_bold=True); i += 1; continue
        if s.startswith('## '):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(5)
            add_runs(p, s[3:], size=13, colour=NAVY, base_bold=True)
            para_rule(p); i += 1; continue
        if s.startswith('# '):
            p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
            add_runs(p, s[2:], size=19, colour=NAVY, base_bold=True); i += 1; continue

        # ---- subtitle ----
        if s.startswith('@@ '):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(10)
            add_runs(p, s[3:], size=9, colour=GREY); i += 1; continue

        # ---- callout ----
        if s.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.35)
            p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
            left_bar(p); shade_para(p, "FAF7F1")
            add_runs(p, s[2:], size=9.2, colour=BLACK); i += 1; continue

        # ---- bullets ----
        m = re.match(r'^(\s*)-\s+(.*)$', ln.rstrip())
        if m:
            depth = len(m.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.45 + 0.45 * depth)
            p.paragraph_format.first_line_indent = Cm(-0.28)
            p.paragraph_format.space_after = Pt(2.5)
            bullet = '\u2022  ' if depth == 0 else '\u2013  '
            br = p.add_run(bullet)
            br.font.size = Pt(8.5); br.font.name = 'Georgia'
            br.font.color.rgb = ACCENT if depth == 0 else GREY
            add_runs(p, m.group(2), size=9.5)
            i += 1; continue

        # ---- numbered ----
        m = re.match(r'^(\s*)(\d+)\.\s+(.*)$', ln.rstrip())
        if m:
            depth = len(m.group(1)) // 2
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.55 + 0.45 * depth)
            p.paragraph_format.first_line_indent = Cm(-0.55)
            p.paragraph_format.space_after = Pt(2.5)
            nr = p.add_run(f'{m.group(2)}.  ')
            nr.font.size = Pt(9.5); nr.font.name = 'Georgia'; nr.bold = True; nr.font.color.rgb = NAVY
            add_runs(p, m.group(3), size=9.5)
            i += 1; continue

        # ---- small print ----
        if s.startswith('~ '):
            p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            add_runs(p, s[2:], size=7.6, colour=GREY); i += 1; continue

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_runs(p, s, size=9.5)
        i += 1

    normalise_order(doc)
    doc.save(out_path)
    print('wrote', out_path)


if __name__ == '__main__':
    if len(sys.argv) > 3:
        FOOTER = sys.argv[3]
    build(sys.argv[1], sys.argv[2])
